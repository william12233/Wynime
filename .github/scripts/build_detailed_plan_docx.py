from __future__ import annotations

import os
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "Wynime_完整應用計畫與技術規格_v0.3.docx"
SOURCES = (
    ("完整產品與技術計畫", ROOT / "docs" / "PROJECT_PLAN.md"),
    ("架構邊界", ROOT / "docs" / "ARCHITECTURE.md"),
    ("已核准技術決策", ROOT / "docs" / "DECISIONS.md"),
    ("Codex／Agent 工作規則", ROOT / "AGENTS.md"),
)


def set_shading(element, fill: str) -> None:
    properties = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def set_keep_with_next(paragraph) -> None:
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\([^)]*\)", r"\1", text)
    return text.replace("**", "").replace("__", "").replace("`", "").strip()


def add_inline(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*.*?\*\*|`.*?`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
        else:
            paragraph.add_run(re.sub(r"\[([^]]+)\]\([^)]*\)", r"\1", part))


def add_table(document: Document, lines: list[str]) -> None:
    rows = [[clean_inline(cell) for cell in line.strip().strip("|").split("|")] for line in lines]
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, values in enumerate(rows):
        for column_index in range(width):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = values[column_index] if column_index < len(values) else ""
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    if row_index == 0:
                        run.bold = True
            if row_index == 0:
                set_shading(cell._tc, "DCEAF2")
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    set_repeat_header(table.rows[0])
    document.add_paragraph()


def add_markdown(document: Document, text: str) -> None:
    lines = text.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Cm(0.6)
                paragraph.paragraph_format.right_indent = Cm(0.3)
                paragraph.paragraph_format.space_after = Pt(6)
                set_shading(paragraph._p, "F3F5F7")
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Consolas")
                run.font.size = Pt(8.5)
                code_lines.clear()
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            paragraph = document.add_heading(clean_inline(heading.group(2)), level=min(len(heading.group(1)), 4))
            set_keep_with_next(paragraph)
            index += 1
            continue
        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet:
            add_inline(document.add_paragraph(style="List Bullet"), bullet.group(1))
            index += 1
            continue
        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if numbered:
            add_inline(document.add_paragraph(style="List Number"), numbered.group(1))
            index += 1
            continue
        quote = re.match(r"^>\s?(.*)$", stripped)
        if quote:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.7)
            paragraph.paragraph_format.right_indent = Cm(0.4)
            set_shading(paragraph._p, "EEF4F8")
            paragraph.add_run(clean_inline(quote.group(1))).italic = True
            index += 1
            continue
        if stripped and stripped != "---":
            add_inline(document.add_paragraph(), stripped)
        index += 1


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Noto Sans CJK TC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK TC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color in (
        ("Title", 29, RGBColor(35, 60, 73)),
        ("Heading 1", 18, RGBColor(35, 60, 73)),
        ("Heading 2", 14, RGBColor(56, 92, 110)),
        ("Heading 3", 12, RGBColor(70, 105, 120)),
        ("Heading 4", 11, RGBColor(80, 105, 115)),
    ):
        style = document.styles[style_name]
        style.font.name = "Noto Sans CJK TC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK TC")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    configure_styles(document)

    properties = document.core_properties
    properties.title = "Wynime 完整應用計畫與技術規格"
    properties.subject = "Wynime Android / Windows 跨平台動畫來源播放器產品與技術計畫"
    properties.author = "Wynime Project"
    properties.keywords = "Wynime, Flutter, Android, Windows, HLS, M3U8, Bangumi, Download"
    properties.comments = "Generated from repository Markdown; Markdown and ADR files remain authoritative."

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    run = title.add_run("Wynime")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor(35, 60, 73)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("完整應用計畫與技術規格").bold = True
    version = document.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.add_run("版本 v0.3｜由 GitHub 權威文件自動產生").font.color.rgb = RGBColor(95, 95, 95)

    notice = document.add_table(rows=1, cols=1)
    notice.style = "Table Grid"
    notice.alignment = WD_TABLE_ALIGNMENT.CENTER
    notice.cell(0, 0).text = (
        "文件治理：本 DOCX 是人工閱讀快照；AGENTS.md、docs/PROJECT_PLAN.md、"
        "docs/ARCHITECTURE.md 與 docs/DECISIONS.md 才是權威來源，有衝突時以 Markdown 為準。"
    )
    set_shading(notice.cell(0, 0)._tc, "EEF4F8")
    for paragraph in notice.cell(0, 0).paragraphs:
        for notice_run in paragraph.runs:
            notice_run.bold = True
            notice_run.font.size = Pt(10)

    document.add_page_break()
    document.add_heading("文件來源與讀取順序", level=1)
    for order, (_, source) in enumerate(SOURCES, start=1):
        document.add_paragraph(f"{order}. {source.relative_to(ROOT).as_posix()}", style="List Number")
    document.add_paragraph(f"來源 Commit：{os.environ.get('SOURCE_COMMIT', 'unknown')}")
    document.add_paragraph(f"產生時間：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")

    for heading, source in SOURCES:
        document.add_page_break()
        document.add_heading(heading, level=1)
        add_markdown(document, source.read_text(encoding="utf-8"))

    document.add_page_break()
    document.add_heading("文件維護規則", level=1)
    for rule in (
        "產品需求或核准決策變更時，先修改 Markdown／ADR，再重新產生 DOCX。",
        "DOCX 不作為 Agent 唯一輸入；新工作階段必須先讀 AGENTS.md 與 Markdown。",
        "下載、刪除、去廣告、PlaybackSession 與來源安全契約不得自行弱化。",
        "未核准項目不得因 DOCX 摘要被誤判為已核准。",
    ):
        document.add_paragraph(rule, style="List Bullet")

    for document_section in document.sections:
        header = document_section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header.add_run("Wynime｜完整應用計畫與技術規格｜v0.3")
        header_run.font.size = Pt(7.5)
        header_run.font.color.rgb = RGBColor(110, 110, 110)
        footer = document_section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("Wynime｜GitHub: william12233/Wynime｜人工審核快照")
        footer_run.font.size = Pt(7.5)
        footer_run.font.color.rgb = RGBColor(110, 110, 110)
    return document


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT)
    if not OUTPUT.is_file() or OUTPUT.stat().st_size == 0:
        raise RuntimeError(f"DOCX generation failed: {OUTPUT}")
    print(f"Generated {OUTPUT.relative_to(ROOT)} ({OUTPUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
